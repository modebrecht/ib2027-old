from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK

OUT='hw/P1.docx'

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ['val','sz','space','color']:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'),'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def set_keep_with_next(p, value=True):
    pPr = p._p.get_or_add_pPr()
    k = pPr.find(qn('w:keepNext'))
    if k is None:
        k = OxmlElement('w:keepNext'); pPr.append(k)
    k.set(qn('w:val'), '1' if value else '0')

def set_keep_together(p, value=True):
    pPr = p._p.get_or_add_pPr()
    k = pPr.find(qn('w:keepLines'))
    if k is None:
        k = OxmlElement('w:keepLines'); pPr.append(k)
    k.set(qn('w:val'), '1' if value else '0')

def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run('Seite ')
    r.font.size = Pt(8); r.font.color.rgb=RGBColor(100,116,139)
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text=' PAGE '
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    r._r.append(fldChar1); r._r.append(instrText); r._r.append(fldChar2)

ACCENT='1D4ED8'; DARK='0F172A'; MUTED='64748B'; LIGHT='EFF6FF'; BORDER='CBD5E1'; GREEN='15803D'

doc=Document()
sec=doc.sections[0]
sec.page_width=Cm(21); sec.page_height=Cm(29.7)
sec.top_margin=Cm(1.25); sec.bottom_margin=Cm(1.25); sec.left_margin=Cm(1.55); sec.right_margin=Cm(1.55)

styles=doc.styles
styles['Normal'].font.name='Arial'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial'); styles['Normal'].font.size=Pt(10)
styles['Normal'].paragraph_format.space_after=Pt(4)
for s in ['Title','Heading 1','Heading 2','Heading 3']:
    styles[s].font.name='Arial'; styles[s]._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial')
styles['Title'].font.size=Pt(23); styles['Title'].font.bold=True; styles['Title'].font.color.rgb=RGBColor.from_string(DARK)
styles['Heading 1'].font.size=Pt(15); styles['Heading 1'].font.bold=True; styles['Heading 1'].font.color.rgb=RGBColor.from_string(DARK); styles['Heading 1'].paragraph_format.space_before=Pt(9); styles['Heading 1'].paragraph_format.space_after=Pt(5)
styles['Heading 2'].font.size=Pt(11.5); styles['Heading 2'].font.bold=True; styles['Heading 2'].font.color.rgb=RGBColor.from_string(ACCENT); styles['Heading 2'].paragraph_format.space_before=Pt(6); styles['Heading 2'].paragraph_format.space_after=Pt(3)

footer=sec.footer
p=footer.paragraphs[0]
p.text='P1 · Übungstest Hardware · '
p.runs[0].font.name='Arial'; p.runs[0].font.size=Pt(8); p.runs[0].font.color.rgb=RGBColor.from_string(MUTED)
add_page_field(p)

band=doc.add_table(rows=1, cols=1); band.alignment=WD_TABLE_ALIGNMENT.CENTER; band.autofit=False; band.columns[0].width=Cm(17.9)
cell=band.cell(0,0); shade(cell, DARK); set_cell_margins(cell,top=190,bottom=190,start=220,end=220)
p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
r=p.add_run('P1  |  ÜBUNGSTEST HARDWARE'); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=RGBColor(255,255,255); r.font.name='Arial'
p2=cell.add_paragraph(); p2.paragraph_format.space_before=Pt(2); p2.paragraph_format.space_after=Pt(0)
r=p2.add_run('Hardware A1–A14  ·  ca. 35–40 Minuten  ·  ohne Kursseiten'); r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(203,213,225); r.font.name='Arial'

doc.add_paragraph().paragraph_format.space_after=Pt(0)

meta=doc.add_table(rows=1, cols=3); meta.alignment=WD_TABLE_ALIGNMENT.CENTER; meta.autofit=False
widths=[8.3,4.2,5.4]
labels=['Vorname / Name','Klasse','Datum']
for i,(w,label) in enumerate(zip(widths,labels)):
    c=meta.cell(0,i); c.width=Cm(w); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(c,top=110,bottom=120,start=120,end=120)
    set_cell_border(c, bottom={'val':'single','sz':'8','color':BORDER})
    p=c.paragraphs[0]; r=p.add_run(label.upper()); r.font.size=Pt(7.5); r.bold=True; r.font.color.rgb=RGBColor.from_string(MUTED)
    p2=c.add_paragraph('____________________________'); p2.paragraph_format.space_before=Pt(5); p2.paragraph_format.space_after=Pt(0); p2.runs[0].font.size=Pt(8); p2.runs[0].font.color.rgb=RGBColor.from_string('94A3B8')

score=doc.add_table(rows=1, cols=3); score.alignment=WD_TABLE_ALIGNMENT.CENTER; score.autofit=False
score_widths=[8.3,4.2,5.4]
score_labels=[('ERREICHTE PUNKTE','_____ / 41'),('GESAMTPUNKTE','41'),('NOTE','_____')]
for i,(w,(lab,val)) in enumerate(zip(score_widths,score_labels)):
    c=score.cell(0,i); c.width=Cm(w); shade(c, LIGHT if i!=2 else 'ECFDF5'); set_cell_margins(c,top=120,bottom=120,start=120,end=120)
    set_cell_border(c, top={'val':'single','sz':'5','color':'BFDBFE'}, bottom={'val':'single','sz':'5','color':'BFDBFE'}, left={'val':'single','sz':'5','color':'BFDBFE'}, right={'val':'single','sz':'5','color':'BFDBFE'})
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(lab+'\n'); r.font.size=Pt(7.5); r.bold=True; r.font.color.rgb=RGBColor.from_string(MUTED)
    r2=p.add_run(val); r2.font.size=Pt(15); r2.bold=True; r2.font.color.rgb=RGBColor.from_string(GREEN if i==2 else DARK)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(7); p.paragraph_format.space_after=Pt(3)
r=p.add_run('NOTENSCHLÜSSEL'); r.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MUTED)
key=doc.add_table(rows=2, cols=6); key.alignment=WD_TABLE_ALIGNMENT.CENTER; key.autofit=False
points=['41 P','33 P','25 P','17 P','9 P','0 P']; grades=['6.0','5.0','4.0','3.0','2.0','1.0']
for j in range(6):
    for i in range(2):
        c=key.cell(i,j); c.width=Cm(2.95); set_cell_margins(c,top=65,bottom=65,start=50,end=50); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(c, top={'val':'single','sz':'4','color':BORDER}, bottom={'val':'single','sz':'4','color':BORDER}, left={'val':'single','sz':'4','color':BORDER}, right={'val':'single','sz':'4','color':BORDER})
        if i==0: shade(c,'F8FAFC')
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(points[j] if i==0 else grades[j]); r.font.size=Pt(8.5); r.bold=(i==1); r.font.color.rgb=RGBColor.from_string(DARK if i==1 else MUTED)
form=doc.add_paragraph(); form.alignment=WD_ALIGN_PARAGRAPH.RIGHT; form.paragraph_format.space_before=Pt(2); form.paragraph_format.space_after=Pt(7)
r=form.add_run('Linear: Note = 1 + 5 × (erreichte Punkte / 41)  ·  60 % = Note 4.0'); r.font.size=Pt(7.5); r.font.color.rgb=RGBColor.from_string(MUTED)

def add_task_heading(num,title,points):
    p=doc.add_paragraph(style='Heading 1'); set_keep_with_next(p,True)
    r=p.add_run(f'{num}. {title}'); r.font.color.rgb=RGBColor.from_string(DARK)
    rr=p.add_run(f'   {points} P'); rr.font.size=Pt(10); rr.font.color.rgb=RGBColor.from_string(ACCENT)
    return p

def add_instruction(text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5); set_keep_with_next(p,True)
    r=p.add_run(text); r.font.size=Pt(9.5); r.font.color.rgb=RGBColor.from_string(MUTED)

def add_answer_lines(n=2):
    for _ in range(n):
        p=doc.add_paragraph('________________________________________________________________________________')
        p.paragraph_format.space_after=Pt(3); p.runs[0].font.size=Pt(9); p.runs[0].font.color.rgb=RGBColor.from_string('94A3B8')

def add_mc_component(name, options):
    p=doc.add_paragraph(style='Heading 2'); set_keep_with_next(p,True); p.add_run(name)
    for opt in options:
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(.15); p.paragraph_format.space_after=Pt(1.5); set_keep_together(p,True)
        r=p.add_run('☐  '+opt); r.font.size=Pt(9.2); r.font.name='Arial'

add_task_heading(1,'Computer-Komponenten',7)
add_instruction('Kreuze bei jeder Komponente die richtige Aussage an. Genau eine Aussage ist korrekt.')
components=[
('CPU',['Speichert Dateien auch ohne Strom dauerhaft.','Führt Berechnungen aus und verarbeitet Befehle von Programmen.','Liefert dem Computer die benötigte elektrische Energie.','Ist hauptsächlich dafür zuständig, Bilder auf dem Bildschirm darzustellen.']),
('RAM',['Speichert Fotos und Dokumente dauerhaft.','Verbindet die verschiedenen Hardware-Komponenten miteinander.','Hält aktuell benötigte Daten und Programme kurzfristig bereit.','Wandelt den Strom aus der Steckdose für die PC-Komponenten um.']),
('SSD',['Speichert Daten dauerhaft und verwendet keine beweglichen mechanischen Teile.','Führt die eigentlichen Berechnungen eines Programms aus.','Speichert Daten nur solange der Computer eingeschaltet ist.','Ist eine langsamere Variante des Arbeitsspeichers.']),
('HDD',['Berechnet 3D-Grafiken für Spiele.','Ist ein dauerhafter Speicher mit mechanisch beweglichen Teilen.','Ist grundsätzlich schneller als eine SSD.','Wird als kurzfristiger Arbeitsspeicher für geöffnete Programme verwendet.']),
('Mainboard',['Speichert das Betriebssystem dauerhaft.','Versorgt sämtliche Komponenten direkt mit Strom aus der Steckdose.','Kühlt Prozessor und Grafikkarte.','Verbindet zentrale Komponenten und ermöglicht deren Kommunikation untereinander.']),
('Netzteil',['Wandelt Netzstrom um und versorgt die PC-Komponenten mit geeigneter elektrischer Energie.','Speichert besonders grosse Dateien günstig.','Verbindet CPU, RAM und Grafikkarte miteinander.','Beschleunigt Programme, indem es häufig benötigte Daten zwischenspeichert.']),
('Grafikkarte',['Speichert Dokumente dauerhaft auf dem Computer.','Berechnet vor allem Grafiken und Bildinformationen für die Ausgabe.','Verteilt den Strom an Mainboard und Laufwerke.','Ist für die allgemeine Verarbeitung aller Programmbefehle verantwortlich.'])]
for name,opts in components: add_mc_component(name,opts)

add_task_heading(2,'EVA-Prinzip',6)
add_instruction('Ordne jeweils Eingabe, Verarbeitung oder Ausgabe zu.')
eva=['Du bewegst die Maus.','Der Computer berechnet die neue Position des Mauszeigers.','Der Mauszeiger erscheint an einer neuen Stelle auf dem Bildschirm.','Ein Barcode wird an der Kasse eingescannt.','Das Kassensystem berechnet den Gesamtpreis.','Der Preis erscheint auf dem Display.']
t=doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Cm(14.4); t.columns[1].width=Cm(3.4)
for i,h in enumerate(['Situation','Antwort (E / V / A)']):
    c=t.cell(0,i); shade(c,DARK); set_cell_margins(c,top=90,bottom=90,start=90,end=90); p=c.paragraphs[0]; r=p.add_run(h); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(255,255,255)
set_repeat_table_header(t.rows[0])
for idx,text in enumerate(eva,1):
    cells=t.add_row().cells
    for c in cells: set_cell_margins(c,top=90,bottom=90,start=90,end=90); set_cell_border(c, bottom={'val':'single','sz':'4','color':BORDER})
    p=cells[0].paragraphs[0]; p.add_run(f'{idx}. {text}').font.size=Pt(9.2)
    p=cells[1].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('________').font.size=Pt(10)

add_task_heading(3,'EVA verstehen',3)
add_instruction('Ein Schüler sagt: „Der Monitor verarbeitet die Daten, weil er das Bild erzeugt.“ Erkläre, warum diese Aussage nicht stimmt.')
add_answer_lines(4)

add_task_heading(4,'Anschlüsse',6)
add_instruction('Wähle den passendsten Anschluss aus der Auswahl und schreibe ihn hin.')
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5)
r=p.add_run('Auswahl: HDMI · RJ45/Ethernet · 3,5-mm-Klinke · USB-C · VGA · M.2'); r.bold=True; r.font.size=Pt(9)
ports=['Moderner Monitor oder Fernseher mit Bild und Ton','Netzwerkkabel zum Router','Kopfhörer mit klassischem Audiostecker','Moderner USB-Stick oder Smartphone','Älterer Monitor mit analogem Bildsignal','Interne SSD direkt auf dem Mainboard']
t=doc.add_table(rows=0, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Cm(14.2); t.columns[1].width=Cm(3.6)
for idx,text in enumerate(ports,1):
    cells=t.add_row().cells
    for c in cells: set_cell_margins(c,top=75,bottom=75,start=90,end=90); set_cell_border(c,bottom={'val':'single','sz':'4','color':BORDER})
    cells[0].paragraphs[0].add_run(f'{idx}. {text}').font.size=Pt(9.2)
    cells[1].paragraphs[0].add_run('________________').font.size=Pt(9)

add_task_heading(5,'RAM, HDD oder SSD?',5)
add_instruction('Wähle jeweils die beste Lösung. Bei den offenen Fragen antworte kurz und präzise.')
items=[
('1. Ein PC soll schneller starten und Programme schneller öffnen.','☐ RAM    ☐ HDD    ☐ SSD'),
('2. 4 TB Fotos sollen möglichst günstig archiviert werden.','☐ RAM    ☐ HDD    ☐ SSD'),
('3. Ein Schüler öffnet gleichzeitig Browser, Teams, PowerPoint und ein Bildbearbeitungsprogramm. Der PC wird dabei langsam.','☐ RAM    ☐ HDD    ☐ SSD')]
for q,opts in items:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1); p.add_run(q).font.size=Pt(9.3)
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(.35); p.paragraph_format.space_after=Pt(4); r=p.add_run(opts); r.font.size=Pt(9.2); r.bold=True
p=doc.add_paragraph('4. Was passiert mit den Daten im RAM, wenn der Strom ausgeschaltet wird?'); p.runs[0].font.size=Pt(9.3); add_answer_lines(1)
p=doc.add_paragraph('5. Nenne einen Vorteil einer SSD gegenüber einer HDD.'); p.runs[0].font.size=Pt(9.3); add_answer_lines(1)

add_task_heading(6,'Kaufberatung',4)
add_instruction('Eine Kundin braucht einen Laptop für Word/PowerPoint, viele Browser-Tabs, Videokonferenzen und Schule. Anspruchsvolle Games spielt sie nicht.')
lt=doc.add_table(rows=2, cols=2); lt.alignment=WD_TABLE_ALIGNMENT.CENTER; lt.autofit=False; lt.columns[0].width=Cm(8.7); lt.columns[1].width=Cm(8.7)
for j,title in enumerate(['LAPTOP A','LAPTOP B']):
    c=lt.cell(0,j); shade(c, ACCENT if j==0 else GREEN); set_cell_margins(c,top=90,bottom=90,start=100,end=100); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(title); r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(11)
vals=['8 GB RAM\n256 GB SSD\nsehr starke Grafikkarte\nCHF 1\'500','16 GB RAM\n512 GB SSD\nnormale integrierte Grafik\nCHF 850']
for j,text in enumerate(vals):
    c=lt.cell(1,j); shade(c,'F8FAFC'); set_cell_margins(c,top=120,bottom=120,start=140,end=140); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(text); r.font.size=Pt(9.5); r.bold=True
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.add_run('Welchen Laptop würdest du empfehlen? Begründe mit mindestens zwei sinnvollen Argumenten.').bold=True
add_answer_lines(5)

add_task_heading(7,'Troubleshooting',6)
add_instruction('Bei jedem Fall gehören zwei Aktionen nicht zu einem sinnvollen ersten Troubleshooting-Ablauf. Streiche diese zwei Aktionen.')
troubles=[
('A) PC startet nicht – 2 P',['Steckdose prüfen','Stromkabel prüfen','Netzschalter am Netzteil prüfen','Windows neu installieren','Monitor heller stellen','andere Steckdose testen']),
('B) Monitor zeigt kein Bild – 2 P',['Prüfen, ob der Monitor eingeschaltet ist','HDMI-/DisplayPort-Kabel prüfen','richtigen Eingang am Monitor auswählen','Tastatur austauschen','Druckertreiber installieren','Kabel an einem anderen Anschluss testen']),
('C) Kein Internet über Netzwerkkabel – 2 P',['RJ45-Kabel prüfen','prüfen, ob das Kabel am Router/Switch steckt','anderen Netzwerkanschluss testen','Grafikkarte ausbauen','Audiokabel ersetzen','Router bzw. Netzwerkstatus prüfen'])]
for title,opts in troubles:
    p=doc.add_paragraph(style='Heading 2'); p.add_run(title); set_keep_with_next(p,True)
    tb=doc.add_table(rows=0, cols=2); tb.alignment=WD_TABLE_ALIGNMENT.LEFT; tb.autofit=False; tb.columns[0].width=Cm(8.9); tb.columns[1].width=Cm(8.9)
    for k in range(0,6,2):
        cells=tb.add_row().cells
        for j in [0,1]:
            c=cells[j]; set_cell_margins(c,top=55,bottom=55,start=60,end=60); set_cell_border(c,bottom={'val':'single','sz':'3','color':'E2E8F0'})
            p=c.paragraphs[0]; p.add_run('□  '+opts[k+j]).font.size=Pt(8.8)

p=doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)
add_task_heading(8,'Green IT',4)
add_instruction('Beantworte die Fragen zur nachhaltigen Nutzung von Geräten.')
p=doc.add_paragraph(); p.add_run('1. Ein funktionierender Laptop ist vier Jahre alt. Der Akku hält nicht mehr lange. Was ist nachhaltiger?').bold=True
p=doc.add_paragraph('☐ A: Neues Gerät kaufen      ☐ B: Akku ersetzen und Laptop weiterverwenden'); p.paragraph_format.left_indent=Cm(.25); p.runs[0].font.size=Pt(9.2)
p=doc.add_paragraph('Begründe deine Entscheidung. (2 P)'); p.runs[0].font.size=Pt(9.2); add_answer_lines(5)
p=doc.add_paragraph('2. Warum ist die Herstellung von Elektronik ökologisch relevant? Nenne einen Grund. (1 P)'); p.runs[0].font.size=Pt(9.2); add_answer_lines(3)
p=doc.add_paragraph('3. Was sollte mit einem defekten Smartphone gemacht werden? (1 P)'); p.runs[0].font.size=Pt(9.2)
p=doc.add_paragraph('☐ A: Hausmüll      ☐ B: Schublade      ☐ C: Fachgerecht zurückgeben/recyceln'); p.paragraph_format.left_indent=Cm(.25); p.runs[0].font.size=Pt(9.2)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Viel Erfolg!'); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(ACCENT)

doc.core_properties.title='P1 – Übungstest Hardware'
doc.core_properties.subject='Hardware A1–A14'
doc.core_properties.author=''
doc.core_properties.keywords='Hardware, Übungstest, Informatik'

doc.save(OUT)
print(OUT)
