from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PATH = 'hw/P1.docx'
DARK='0F172A'; ACCENT='1D4ED8'; MUTED='64748B'

def insert_before(target, new_p):
    target._p.addprevious(new_p._p)

def make_p(doc, text='', size=9.5, bold=False, color=DARK, before=0, after=4, left=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if left:
        p.paragraph_format.left_indent = Cm(left)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    return p

def line(doc):
    p = make_p(doc, '________________________________________________________________________________', size=9, color='94A3B8', after=3)
    return p

doc = Document(PATH)

# Avoid duplicate insertion if the postprocessor is run twice.
if any('Bonus 1 – Hardware Detective' in p.text for p in doc.paragraphs):
    print('Bonus already present')
    raise SystemExit(0)

target = next((p for p in doc.paragraphs if p.text.strip() == 'Viel Erfolg!'), None)
if target is None:
    target = doc.add_paragraph()

items = []
items.append(make_p(doc, 'Bonus – max. +2 P', size=15, bold=True, before=10, after=5))
items.append(make_p(doc, 'Zwei freiwillige Transferfragen. Jede richtige Lösung gibt +1 Bonuspunkt.', size=9.5, color=MUTED, after=6))

items.append(make_p(doc, 'Bonus 1 – Hardware Detective (+1 P)', size=11.5, bold=True, color=ACCENT, before=4, after=3))
items.append(make_p(doc, 'Ein Schüler sagt: „Mein PC ist langsam, also brauche ich sicher eine neue Grafikkarte.“', size=9.3, after=3))
items.append(make_p(doc, 'Der PC hat: 4 GB RAM · HDD · CPU 35 % · GPU 8 % · 17 Browser-Tabs geöffnet', size=9.2, color=MUTED, after=4))
items.append(make_p(doc, 'Was würdest du zuerst ändern – und warum?', size=9.3, bold=True, after=3))
items.append(line(doc)); items.append(line(doc))

items.append(make_p(doc, 'Bonus 2 – Flaschenhals (+1 P)', size=11.5, bold=True, color=ACCENT, before=5, after=3))
items.append(make_p(doc, 'Ein PC hat: 32 GB RAM · starke CPU · starke Grafikkarte · 5400-rpm-HDD · Windows startet in 2 Minuten', size=9.2, after=4))
items.append(make_p(doc, 'Welche Komponente ist der wahrscheinlichste Flaschenhals? Begründe kurz.', size=9.3, bold=True, after=3))
items.append(line(doc)); items.append(line(doc))
items.append(make_p(doc, 'Bonuspunkte zählen zusätzlich zu den 41 regulären Punkten. Maximalnote: 6.0.', size=8.2, color=MUTED, before=3, after=7))

for p in items:
    insert_before(target, p)

doc.save(PATH)
print(PATH)
